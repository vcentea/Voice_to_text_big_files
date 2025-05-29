#!/usr/bin/env python3
"""
Meeting Report Generator - Phase 1: Chunking and Initial Processing

This script implements the first phase of hierarchical map-reduce summarization
for generating meeting summaries, minutes, and action items from long transcripts.

Phase 1 Features:
- Smart chunking of SRT segments while respecting LLM context window limits
- Chunk-level summarization with the local LLM
- Action item identification within each chunk
- Keyword/topic extraction for each chunk
- Structured data storage for subsequent phases

Usage: python generate_report.py <input_srt_file> [output_json_file]
"""

import os
import sys
import srt
import json
import time
from openai import OpenAI
from datetime import datetime
from typing import List, Dict, Any
import re
from docx import Document  # Added for DOCX generation
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- Configuration ---
LLM_API_BASE_URL = "http://localhost:1234/v1"
LLM_MODEL_NAME = "mistral-small-3.1-24b-instruct-2503"
LLM_API_KEY = "dummy"

# Token limits and chunking parameters
MAX_CONTEXT_TOKENS = 8549  # Actual LLM context window (from error message)
PROMPT_OVERHEAD_TOKENS = 3500  # Reserve for prompt and response (increased)
MAX_CHUNK_TOKENS = MAX_CONTEXT_TOKENS - PROMPT_OVERHEAD_TOKENS  # ~5000 tokens per chunk
AVERAGE_TOKENS_PER_CHAR = 0.4  # More conservative approximation: 1 token ≈ 2.5 characters

def estimate_tokens(text: str) -> int:
    """
    Rough estimation of token count for text.
    More accurate would be to use tiktoken, but this approximation works for chunking.
    """
    return int(len(text) * AVERAGE_TOKENS_PER_CHAR)

def load_srt_file(srt_path: str) -> List[srt.Subtitle]:
    """Loads an SRT file and returns a list of subtitle objects."""
    print(f"🔄 Loading SRT file: {srt_path}")
    try:
        with open(srt_path, 'r', encoding='utf-8') as f:
            subs = list(srt.parse(f.read()))
        print(f"✅ SRT file loaded: {len(subs)} segments")
        return subs
    except Exception as e:
        print(f"❌ Error loading SRT file {srt_path}: {e}")
        sys.exit(1)

def create_smart_chunks(segments: List[srt.Subtitle], max_tokens_per_chunk: int) -> List[Dict[str, Any]]:
    """
    Creates semantically coherent chunks from SRT segments.
    Each chunk will contain multiple segments but stay within token limits.
    
    Returns a list of chunk dictionaries with metadata.
    """
    chunks = []
    current_chunk_segments = []
    current_chunk_text = ""
    current_tokens = 0
    
    for i, segment in enumerate(segments):
        segment_text = segment.content
        segment_tokens = estimate_tokens(segment_text)
        
        # Check if adding this segment would exceed the token limit
        if current_tokens + segment_tokens > max_tokens_per_chunk and current_chunk_segments:
            # Finalize current chunk
            chunks.append({
                'chunk_id': len(chunks) + 1,
                'segments': current_chunk_segments.copy(),
                'text': current_chunk_text.strip(),
                'token_count': current_tokens,
                'start_time': current_chunk_segments[0].start,
                'end_time': current_chunk_segments[-1].end,
                'segment_indices': [seg.index for seg in current_chunk_segments]
            })
            
            # Start new chunk
            current_chunk_segments = []
            current_chunk_text = ""
            current_tokens = 0
        
        # Add segment to current chunk
        current_chunk_segments.append(segment)
        current_chunk_text += f"\n{segment_text}"
        current_tokens += segment_tokens
    
    # Don't forget the last chunk
    if current_chunk_segments:
        chunks.append({
            'chunk_id': len(chunks) + 1,
            'segments': current_chunk_segments.copy(),
            'text': current_chunk_text.strip(),
            'token_count': current_tokens,
            'start_time': current_chunk_segments[0].start,
            'end_time': current_chunk_segments[-1].end,
            'segment_indices': [seg.index for seg in current_chunk_segments]
        })
    
    return chunks

def build_chunk_analysis_prompt(chunk_text: str, chunk_id: int) -> str:
    """
    Builds the prompt for analyzing a single chunk.
    Asks for summary, action items, and key topics.
    """
    prompt = f"""You are analyzing a section of a meeting transcript. Please provide a comprehensive analysis of this chunk.

TRANSCRIPT CHUNK {chunk_id}:
{chunk_text}

Please provide your analysis in the following JSON format (ensure valid JSON syntax):

{{
  "summary": "A concise 2-3 sentence summary of the main points discussed in this chunk",
  "action_items": [
    {{
      "description": "Description of the action item",
      "assignee": "Person responsible (if mentioned/identifiable, otherwise 'Not specified')",
      "deadline": "Deadline if mentioned, otherwise 'Not specified'",
      "context": "Brief context about this action item"
    }}
  ],
  "key_topics": ["topic1", "topic2", "topic3"],
  "decisions_made": ["decision1", "decision2"],
  "questions_raised": ["question1", "question2"],
  "speakers_mentioned": ["speaker1", "speaker2"]
}}

Guidelines:
- Be specific and actionable for action items
- Keep summaries concise but informative
- Extract only clear, actionable items as action items
- Key topics should be 1-3 words each
- Include speaker names/roles if clearly identifiable
- If a section is unclear or no action items exist, use empty arrays []

Respond ONLY with the JSON object, no additional text."""

    return prompt

def query_llm_for_chunk_analysis(client: OpenAI, prompt: str, model_name: str, chunk_id: int) -> Dict[str, Any] | None:
    """
    Sends a chunk analysis prompt to the LLM and parses the JSON response.
    """
    try:
        print(f"  💭 Analyzing chunk {chunk_id} with LLM...")
        completion = client.chat.completions.create(
            model=model_name,
            messages=[
                {
                    "role": "system", 
                    "content": "You are an expert meeting analyst. You extract summaries, action items, and key information from meeting transcripts. Always respond with valid JSON format."
                },
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,  # Lower temperature for more consistent structured output
        )
        
        raw_response = completion.choices[0].message.content.strip()
        
        # Try to parse JSON response
        try:
            # Remove any potential markdown code block markers
            if raw_response.startswith("```json"):
                raw_response = raw_response[7:]
            if raw_response.endswith("```"):
                raw_response = raw_response[:-3]
            raw_response = raw_response.strip()
            
            analysis = json.loads(raw_response)
            
            # Validate expected keys exist
            required_keys = ["summary", "action_items", "key_topics", "decisions_made", "questions_raised", "speakers_mentioned"]
            for key in required_keys:
                if key not in analysis:
                    analysis[key] = [] if key != "summary" else "No summary provided"
            
            print(f"  ✅ Chunk {chunk_id} analyzed successfully")
            return analysis
            
        except json.JSONDecodeError as e:
            print(f"  ⚠️ Failed to parse JSON response for chunk {chunk_id}: {e}")
            print(f"    Raw response: {raw_response[:200]}...")
            return None
            
    except Exception as e:
        print(f"  ❌ Error analyzing chunk {chunk_id}: {e}")
        return None

def process_chunks_phase1(chunks: List[Dict[str, Any]], client: OpenAI, model_name: str) -> List[Dict[str, Any]]:
    """
    Process all chunks through Phase 1 analysis.
    Returns a list of processed chunks with analysis results.
    """
    processed_chunks = []
    
    print(f"\n🔍 Phase 1: Processing {len(chunks)} chunks...")
    
    for chunk in chunks:
        chunk_id = chunk['chunk_id']
        chunk_text = chunk['text']
        
        print(f"\nProcessing chunk {chunk_id}/{len(chunks)}")
        print(f"  📊 Chunk spans: {chunk['start_time']} --> {chunk['end_time']}")
        print(f"  📏 Estimated tokens: {chunk['token_count']}")
        print(f"  📝 Segments: {len(chunk['segments'])}")
        
        # Build prompt and query LLM
        prompt = build_chunk_analysis_prompt(chunk_text, chunk_id)
        analysis = query_llm_for_chunk_analysis(client, prompt, model_name, chunk_id)
        
        # Create processed chunk with analysis
        processed_chunk = {
            'chunk_id': chunk_id,
            'metadata': {
                'start_time': str(chunk['start_time']),
                'end_time': str(chunk['end_time']),
                'token_count': chunk['token_count'],
                'segment_count': len(chunk['segments']),
                'segment_indices': chunk['segment_indices']
            },
            'original_text': chunk_text,
            'analysis': analysis if analysis else {
                'summary': f"Analysis failed for chunk {chunk_id}",
                'action_items': [],
                'key_topics': [],
                'decisions_made': [],
                'questions_raised': [],
                'speakers_mentioned': []
            },
            'processing_timestamp': datetime.now().isoformat()
        }
        
        processed_chunks.append(processed_chunk)
        
        # Small delay to avoid overwhelming the LLM
        time.sleep(0.5)
    
    return processed_chunks

def save_phase1_results(processed_chunks: List[Dict[str, Any]], output_path: str) -> None:
    """
    Save Phase 1 results to a JSON file for use in subsequent phases.
    """
    results = {
        'phase': 1,
        'processing_timestamp': datetime.now().isoformat(),
        'total_chunks': len(processed_chunks),
        'chunks': processed_chunks,
        'statistics': {
            'total_action_items': sum(len(chunk['analysis']['action_items']) for chunk in processed_chunks),
            'total_decisions': sum(len(chunk['analysis']['decisions_made']) for chunk in processed_chunks),
            'total_questions': sum(len(chunk['analysis']['questions_raised']) for chunk in processed_chunks),
            'total_segments_processed': sum(chunk['metadata']['segment_count'] for chunk in processed_chunks)
        }
    }
    
    try:
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(results, f, indent=2, ensure_ascii=False)
        print(f"\n✅ Phase 1 results saved to: {output_path}")
        
        # Print summary statistics
        stats = results['statistics']
        print(f"\n📊 Phase 1 Summary:")
        print(f"   Chunks processed: {results['total_chunks']}")
        print(f"   Total segments: {stats['total_segments_processed']}")
        print(f"   Action items found: {stats['total_action_items']}")
        print(f"   Decisions identified: {stats['total_decisions']}")
        print(f"   Questions raised: {stats['total_questions']}")
        
    except Exception as e:
        print(f"❌ Error saving Phase 1 results: {e}")

def generate_phase1_docx(processed_chunks: List[Dict[str, Any]], output_path: str) -> None:
    """
    Generate a professionally formatted DOCX document with Phase 1 analysis results.
    """
    try:
        doc = Document()
        
        # Title and header
        title = doc.add_heading('Meeting Analysis Report - Phase 1', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata section
        doc.add_heading('Document Information', level=1)
        metadata_para = doc.add_paragraph()
        metadata_para.add_run('Generated: ').bold = True
        metadata_para.add_run(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        metadata_para.add_run('Analysis Phase: ').bold = True
        metadata_para.add_run('1 (Initial Chunking and Processing)\n')
        metadata_para.add_run('Total Chunks Processed: ').bold = True
        metadata_para.add_run(f"{len(processed_chunks)}")
        
        # Statistics overview
        doc.add_heading('Executive Summary', level=1)
        stats = {
            'total_action_items': sum(len(chunk['analysis']['action_items']) for chunk in processed_chunks),
            'total_decisions': sum(len(chunk['analysis']['decisions_made']) for chunk in processed_chunks),
            'total_questions': sum(len(chunk['analysis']['questions_raised']) for chunk in processed_chunks),
            'total_segments_processed': sum(chunk['metadata']['segment_count'] for chunk in processed_chunks)
        }
        
        stats_para = doc.add_paragraph()
        stats_para.add_run('Total Segments Analyzed: ').bold = True
        stats_para.add_run(f"{stats['total_segments_processed']}\n")
        stats_para.add_run('Action Items Identified: ').bold = True
        stats_para.add_run(f"{stats['total_action_items']}\n")
        stats_para.add_run('Decisions Made: ').bold = True
        stats_para.add_run(f"{stats['total_decisions']}\n")
        stats_para.add_run('Questions Raised: ').bold = True
        stats_para.add_run(f"{stats['total_questions']}")
        
        # Consolidated action items section
        doc.add_heading('All Action Items', level=1)
        all_action_items = []
        for chunk in processed_chunks:
            for item in chunk['analysis']['action_items']:
                all_action_items.append({
                    'chunk': chunk['chunk_id'],
                    'time_range': f"{chunk['metadata']['start_time']} - {chunk['metadata']['end_time']}",
                    **item
                })
        
        if all_action_items:
            for i, item in enumerate(all_action_items, 1):
                item_para = doc.add_paragraph()
                item_para.add_run(f"{i}. ").bold = True
                item_para.add_run(f"{item['description']}\n")
                item_para.add_run('   Assignee: ').italic = True
                item_para.add_run(f"{item['assignee']}\n")
                item_para.add_run('   Deadline: ').italic = True
                item_para.add_run(f"{item['deadline']}\n")
                item_para.add_run('   Context: ').italic = True
                item_para.add_run(f"{item['context']}\n")
                item_para.add_run('   Time Range: ').italic = True
                item_para.add_run(f"{item['time_range']}")
        else:
            doc.add_paragraph("No action items were identified in this analysis.")
        
        # Consolidated decisions section
        doc.add_heading('All Decisions Made', level=1)
        all_decisions = []
        for chunk in processed_chunks:
            for decision in chunk['analysis']['decisions_made']:
                all_decisions.append({
                    'decision': decision,
                    'chunk': chunk['chunk_id'],
                    'time_range': f"{chunk['metadata']['start_time']} - {chunk['metadata']['end_time']}"
                })
        
        if all_decisions:
            for i, decision in enumerate(all_decisions, 1):
                decision_para = doc.add_paragraph()
                decision_para.add_run(f"{i}. ").bold = True
                decision_para.add_run(f"{decision['decision']}")
                decision_para.add_run(f" (Time: {decision['time_range']})")
        else:
            doc.add_paragraph("No specific decisions were identified in this analysis.")
        
        # Consolidated questions section
        doc.add_heading('Outstanding Questions', level=1)
        all_questions = []
        for chunk in processed_chunks:
            for question in chunk['analysis']['questions_raised']:
                all_questions.append({
                    'question': question,
                    'chunk': chunk['chunk_id'],
                    'time_range': f"{chunk['metadata']['start_time']} - {chunk['metadata']['end_time']}"
                })
        
        if all_questions:
            for i, question in enumerate(all_questions, 1):
                question_para = doc.add_paragraph()
                question_para.add_run(f"{i}. ").bold = True
                question_para.add_run(f"{question['question']}")
                question_para.add_run(f" (Time: {question['time_range']})")
        else:
            doc.add_paragraph("No outstanding questions were identified in this analysis.")
        
        # Detailed chunk analysis
        doc.add_heading('Detailed Chunk Analysis', level=1)
        
        for chunk in processed_chunks:
            # Chunk header
            chunk_heading = doc.add_heading(f"Chunk {chunk['chunk_id']}", level=2)
            
            # Chunk metadata
            meta_para = doc.add_paragraph()
            meta_para.add_run('Time Range: ').bold = True
            meta_para.add_run(f"{chunk['metadata']['start_time']} - {chunk['metadata']['end_time']}\n")
            meta_para.add_run('Segments: ').bold = True
            meta_para.add_run(f"{chunk['metadata']['segment_count']}\n")
            meta_para.add_run('Estimated Tokens: ').bold = True
            meta_para.add_run(f"{chunk['metadata']['token_count']}")
            
            # Summary
            doc.add_heading('Summary', level=3)
            doc.add_paragraph(chunk['analysis']['summary'])
            
            # Key topics
            if chunk['analysis']['key_topics']:
                doc.add_heading('Key Topics', level=3)
                topics_para = doc.add_paragraph()
                topics_para.add_run(', '.join(chunk['analysis']['key_topics']))
            
            # Speakers mentioned
            if chunk['analysis']['speakers_mentioned']:
                doc.add_heading('Speakers Mentioned', level=3)
                speakers_para = doc.add_paragraph()
                speakers_para.add_run(', '.join(chunk['analysis']['speakers_mentioned']))
            
            # Add some space between chunks
            doc.add_paragraph()
        
        # Save the document
        doc.save(output_path)
        print(f"✅ Phase 1 DOCX report saved to: {output_path}")
        
    except Exception as e:
        print(f"❌ Error generating Phase 1 DOCX report: {e}")

def main():
    if len(sys.argv) < 2:
        print("Usage: python generate_report.py <input_srt_file> [output_json_file]")
        print("Example: python generate_report.py output2_latest.txt phase1_analysis.json")
        sys.exit(1)
    
    input_srt_path = sys.argv[1]
    
    if not os.path.exists(input_srt_path):
        print(f"❌ Input SRT file not found: {input_srt_path}")
        sys.exit(1)
    
    # Determine output path
    if len(sys.argv) > 2:
        output_json_path = sys.argv[2]
    else:
        base, _ = os.path.splitext(input_srt_path)
        output_json_path = f"{base}_phase1_analysis.json"
    
    # Determine DOCX output path
    docx_output_path = output_json_path.replace('.json', '.docx')
    
    print("🚀 Meeting Report Generator - Phase 1")
    print(f"   Input SRT: {input_srt_path}")
    print(f"   Output JSON: {output_json_path}")
    print(f"   Output DOCX: {docx_output_path}")
    print(f"   LLM Endpoint: {LLM_API_BASE_URL}")
    print(f"   LLM Model: {LLM_MODEL_NAME}")
    print(f"   Max tokens per chunk: {MAX_CHUNK_TOKENS}")
    
    # Initialize LLM client
    try:
        client = OpenAI(base_url=LLM_API_BASE_URL, api_key=LLM_API_KEY)
        print("✅ LLM client initialized")
    except Exception as e:
        print(f"❌ Failed to initialize LLM client: {e}")
        sys.exit(1)
    
    # Load and process SRT file
    segments = load_srt_file(input_srt_path)
    
    # Create smart chunks
    print(f"\n📦 Creating smart chunks (max {MAX_CHUNK_TOKENS} tokens each)...")
    chunks = create_smart_chunks(segments, MAX_CHUNK_TOKENS)
    print(f"✅ Created {len(chunks)} chunks from {len(segments)} segments")
    
    # Process chunks through Phase 1
    processed_chunks = process_chunks_phase1(chunks, client, LLM_MODEL_NAME)
    
    # Save results
    save_phase1_results(processed_chunks, output_json_path)
    
    # Generate DOCX report
    generate_phase1_docx(processed_chunks, docx_output_path)
    
    print(f"\n🏁 Phase 1 complete!")
    print(f"   📊 JSON results saved to: {output_json_path}")
    print(f"   📄 DOCX report saved to: {docx_output_path}")
    print("   Next: Run Phase 2 to combine chunk summaries into intermediate summaries")
    print("   Then: Run Phase 3 to generate final meeting minutes and action items")

if __name__ == "__main__":
    main() 