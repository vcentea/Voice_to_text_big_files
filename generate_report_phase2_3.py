#!/usr/bin/env python3
"""
Meeting Report Generator - Phase 2 & 3: Intermediate Summarization and Final Report

This script implements Phase 2 (combining chunk summaries into intermediate summaries) 
and Phase 3 (generating final meeting minutes and action items) of the hierarchical 
map-reduce summarization strategy.

Phase 2 Features:
- Groups Phase 1 chunk summaries into meta-chunks
- Creates intermediate summaries with consolidated action items
- Maintains hierarchical structure for final processing

Phase 3 Features:
- Generates executive summary from intermediate summaries
- Creates comprehensive meeting minutes
- Consolidates all action items with priority/urgency analysis
- Produces professional DOCX final report
- Identifies key decisions, outstanding questions, and participants

Usage: python generate_report_phase2_3.py <phase1_json_file> [output_base_name]
"""

import os
import sys
import json
import time
from openai import OpenAI
from datetime import datetime
from typing import List, Dict, Any
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.dml import MSO_THEME_COLOR_INDEX

# --- Configuration ---
LLM_API_BASE_URL = "http://localhost:1234/v1"
LLM_MODEL_NAME = "mistral-small-3.1-24b-instruct-2503"
LLM_API_KEY = "dummy"

# Phase 2 parameters
MAX_CONTEXT_TOKENS = 8549
PROMPT_OVERHEAD_TOKENS = 3500
MAX_META_CHUNK_TOKENS = MAX_CONTEXT_TOKENS - PROMPT_OVERHEAD_TOKENS
AVERAGE_TOKENS_PER_CHAR = 0.4
META_CHUNK_SIZE = 3  # Number of Phase 1 chunks to combine into meta-chunks

def estimate_tokens(text: str) -> int:
    """Rough estimation of token count for text."""
    return int(len(text) * AVERAGE_TOKENS_PER_CHAR)

def load_phase1_results(json_path: str) -> Dict[str, Any]:
    """Load Phase 1 results from JSON file."""
    print(f"🔄 Loading Phase 1 results: {json_path}")
    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        print(f"✅ Phase 1 results loaded: {data['total_chunks']} chunks")
        return data
    except Exception as e:
        print(f"❌ Error loading Phase 1 results {json_path}: {e}")
        sys.exit(1)

def create_meta_chunks(phase1_chunks: List[Dict[str, Any]], meta_chunk_size: int) -> List[Dict[str, Any]]:
    """
    Groups Phase 1 chunks into meta-chunks for Phase 2 processing.
    """
    meta_chunks = []
    
    for i in range(0, len(phase1_chunks), meta_chunk_size):
        chunk_group = phase1_chunks[i:i + meta_chunk_size]
        
        # Combine summaries and metadata
        combined_summaries = []
        combined_action_items = []
        combined_decisions = []
        combined_questions = []
        combined_topics = set()
        combined_speakers = set()
        
        start_time = chunk_group[0]['metadata']['start_time']
        end_time = chunk_group[-1]['metadata']['end_time']
        total_segments = sum(chunk['metadata']['segment_count'] for chunk in chunk_group)
        
        for chunk in chunk_group:
            analysis = chunk['analysis']
            combined_summaries.append(f"Chunk {chunk['chunk_id']}: {analysis['summary']}")
            combined_action_items.extend(analysis['action_items'])
            combined_decisions.extend(analysis['decisions_made'])
            combined_questions.extend(analysis['questions_raised'])
            combined_topics.update(analysis['key_topics'])
            combined_speakers.update(analysis['speakers_mentioned'])
        
        meta_chunk = {
            'meta_chunk_id': len(meta_chunks) + 1,
            'source_chunks': [chunk['chunk_id'] for chunk in chunk_group],
            'metadata': {
                'start_time': start_time,
                'end_time': end_time,
                'total_segments': total_segments,
                'source_chunk_count': len(chunk_group)
            },
            'combined_content': {
                'summaries': combined_summaries,
                'action_items': combined_action_items,
                'decisions_made': combined_decisions,
                'questions_raised': combined_questions,
                'key_topics': list(combined_topics),
                'speakers_mentioned': list(combined_speakers)
            }
        }
        
        meta_chunks.append(meta_chunk)
    
    return meta_chunks

def build_phase2_prompt(meta_chunk: Dict[str, Any]) -> str:
    """Build prompt for Phase 2 intermediate summarization."""
    content = meta_chunk['combined_content']
    
    prompt = f"""You are analyzing a group of {len(meta_chunk['source_chunks'])} meeting segments that have been pre-summarized. Please create an intermediate-level summary that synthesizes the information.

TIME RANGE: {meta_chunk['metadata']['start_time']} - {meta_chunk['metadata']['end_time']}
SEGMENTS COVERED: {meta_chunk['metadata']['total_segments']} segments

INDIVIDUAL CHUNK SUMMARIES:
{chr(10).join(content['summaries'])}

CONSOLIDATED ACTION ITEMS FROM THIS SECTION:
{chr(10).join([f"- {item['description']} (Assignee: {item['assignee']}, Deadline: {item['deadline']})" for item in content['action_items']]) if content['action_items'] else "No action items identified"}

DECISIONS MADE IN THIS SECTION:
{chr(10).join([f"- {decision}" for decision in content['decisions_made']]) if content['decisions_made'] else "No specific decisions identified"}

KEY TOPICS DISCUSSED: {', '.join(content['key_topics']) if content['key_topics'] else "No specific topics identified"}

Please provide a comprehensive intermediate analysis in JSON format:

{{
  "intermediate_summary": "A comprehensive 3-4 sentence summary that synthesizes the main themes and outcomes from this section of the meeting",
  "key_themes": ["theme1", "theme2", "theme3"],
  "consolidated_action_items": [
    {{
      "description": "Refined/consolidated action item description",
      "assignee": "Person responsible",
      "deadline": "Timeline if specified",
      "priority": "High/Medium/Low",
      "context": "Context and background"
    }}
  ],
  "important_decisions": ["decision1", "decision2"],
  "strategic_insights": ["insight1", "insight2"],
  "outstanding_issues": ["issue1", "issue2"],
  "next_steps": ["step1", "step2"]
}}

Guidelines:
- Synthesize rather than summarize - look for connections and patterns
- Prioritize action items by importance and urgency
- Identify strategic implications and insights
- Consolidate similar or related items
- Focus on actionable outcomes and decisions

Respond ONLY with the JSON object."""

    return prompt

def query_llm_phase2(client: OpenAI, prompt: str, model_name: str, meta_chunk_id: int) -> Dict[str, Any] | None:
    """Send Phase 2 prompt to LLM and parse response."""
    try:
        print(f"  💭 Processing meta-chunk {meta_chunk_id} with LLM...")
        completion = client.chat.completions.create(
            model=model_name,
            messages=[
                {
                    "role": "system",
                    "content": "You are an expert meeting analyst specializing in synthesis and strategic analysis. You excel at identifying patterns, consolidating information, and extracting actionable insights from complex discussions. Always respond with valid JSON."
                },
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
        )
        
        raw_response = completion.choices[0].message.content.strip()
        
        # Parse JSON response
        try:
            if raw_response.startswith("```json"):
                raw_response = raw_response[7:]
            if raw_response.endswith("```"):
                raw_response = raw_response[:-3]
            raw_response = raw_response.strip()
            
            analysis = json.loads(raw_response)
            
            # Validate expected keys
            required_keys = ["intermediate_summary", "key_themes", "consolidated_action_items", 
                           "important_decisions", "strategic_insights", "outstanding_issues", "next_steps"]
            for key in required_keys:
                if key not in analysis:
                    analysis[key] = [] if key != "intermediate_summary" else "No summary provided"
            
            print(f"  ✅ Meta-chunk {meta_chunk_id} processed successfully")
            return analysis
            
        except json.JSONDecodeError as e:
            print(f"  ⚠️ Failed to parse JSON response for meta-chunk {meta_chunk_id}: {e}")
            print(f"    Raw response: {raw_response[:200]}...")
            return None
            
    except Exception as e:
        print(f"  ❌ Error processing meta-chunk {meta_chunk_id}: {e}")
        return None

def process_phase2(meta_chunks: List[Dict[str, Any]], client: OpenAI, model_name: str) -> List[Dict[str, Any]]:
    """Process all meta-chunks through Phase 2."""
    processed_meta_chunks = []
    
    print(f"\n🔍 Phase 2: Processing {len(meta_chunks)} meta-chunks...")
    
    for meta_chunk in meta_chunks:
        meta_chunk_id = meta_chunk['meta_chunk_id']
        
        print(f"\nProcessing meta-chunk {meta_chunk_id}/{len(meta_chunks)}")
        print(f"  📊 Time range: {meta_chunk['metadata']['start_time']} --> {meta_chunk['metadata']['end_time']}")
        print(f"  📝 Source chunks: {meta_chunk['source_chunks']}")
        print(f"  🎯 Total segments: {meta_chunk['metadata']['total_segments']}")
        
        prompt = build_phase2_prompt(meta_chunk)
        analysis = query_llm_phase2(client, prompt, model_name, meta_chunk_id)
        
        processed_meta_chunk = {
            'meta_chunk_id': meta_chunk_id,
            'metadata': meta_chunk['metadata'],
            'source_chunks': meta_chunk['source_chunks'],
            'phase2_analysis': analysis if analysis else {
                'intermediate_summary': f"Analysis failed for meta-chunk {meta_chunk_id}",
                'key_themes': [],
                'consolidated_action_items': [],
                'important_decisions': [],
                'strategic_insights': [],
                'outstanding_issues': [],
                'next_steps': []
            },
            'processing_timestamp': datetime.now().isoformat()
        }
        
        processed_meta_chunks.append(processed_meta_chunk)
        time.sleep(0.5)
    
    return processed_meta_chunks

def build_phase3_prompt(phase2_results: List[Dict[str, Any]], original_stats: Dict[str, Any]) -> str:
    """Build prompt for Phase 3 final report generation."""
    
    # Collect all intermediate summaries and insights
    all_summaries = []
    all_themes = set()
    all_action_items = []
    all_decisions = []
    all_insights = []
    all_issues = []
    all_next_steps = []
    
    for meta_chunk in phase2_results:
        analysis = meta_chunk['phase2_analysis']
        all_summaries.append(f"Section {meta_chunk['meta_chunk_id']} ({meta_chunk['metadata']['start_time']} - {meta_chunk['metadata']['end_time']}): {analysis['intermediate_summary']}")
        all_themes.update(analysis['key_themes'])
        all_action_items.extend(analysis['consolidated_action_items'])
        all_decisions.extend(analysis['important_decisions'])
        all_insights.extend(analysis['strategic_insights'])
        all_issues.extend(analysis['outstanding_issues'])
        all_next_steps.extend(analysis['next_steps'])
    
    prompt = f"""You are creating the final comprehensive meeting report from intermediate analyses. This is the culmination of a hierarchical analysis process.

MEETING OVERVIEW:
- Total Segments Processed: {original_stats['total_segments_processed']}
- Original Chunks: Multiple sections analyzed
- Duration: Full meeting session

INTERMEDIATE SECTION SUMMARIES:
{chr(10).join(all_summaries)}

CONSOLIDATED ACTION ITEMS:
{chr(10).join([f"- {item['description']} (Assignee: {item['assignee']}, Priority: {item.get('priority', 'Not specified')}, Deadline: {item['deadline']})" for item in all_action_items]) if all_action_items else "No action items identified"}

KEY THEMES IDENTIFIED: {', '.join(all_themes) if all_themes else "No major themes identified"}

IMPORTANT DECISIONS: {chr(10).join([f"- {decision}" for decision in all_decisions]) if all_decisions else "No major decisions identified"}

STRATEGIC INSIGHTS: {chr(10).join([f"- {insight}" for insight in all_insights]) if all_insights else "No strategic insights identified"}

Please generate a comprehensive final meeting report in JSON format:

{{
  "executive_summary": "A compelling 4-5 sentence executive summary that captures the essence, outcomes, and significance of the entire meeting",
  "meeting_purpose": "Inferred primary purpose/objective of this meeting",
  "key_participants": ["participant1", "participant2"],
  "main_topics_discussed": [
    {{
      "topic": "Topic name",
      "description": "What was discussed about this topic",
      "outcome": "Result or conclusion reached"
    }}
  ],
  "final_action_items": [
    {{
      "id": 1,
      "description": "Clear, actionable item description",
      "assignee": "Person responsible",
      "deadline": "Timeline",
      "priority": "High/Medium/Low",
      "dependencies": "What needs to happen first",
      "success_criteria": "How we know it's done"
    }}
  ],
  "decisions_and_resolutions": [
    {{
      "decision": "What was decided",
      "rationale": "Why this decision was made",
      "impact": "Expected impact or implications"
    }}
  ],
  "strategic_outcomes": ["outcome1", "outcome2"],
  "risks_and_concerns": ["risk1", "concern1"],
  "follow_up_required": ["followup1", "followup2"],
  "recommendations": ["recommendation1", "recommendation2"],
  "meeting_effectiveness_assessment": "Brief assessment of how productive the meeting was"
}}

Guidelines:
- Focus on strategic significance and business impact
- Ensure action items are specific, measurable, and actionable
- Identify patterns and themes across the entire discussion
- Highlight the most important outcomes and decisions
- Provide executive-level insights suitable for leadership review

Respond ONLY with the JSON object."""

    return prompt

def query_llm_phase3(client: OpenAI, prompt: str, model_name: str) -> Dict[str, Any] | None:
    """Send Phase 3 prompt to LLM and parse response."""
    try:
        print(f"  💭 Generating final meeting report with LLM...")
        completion = client.chat.completions.create(
            model=model_name,
            messages=[
                {
                    "role": "system",
                    "content": "You are a senior executive assistant and meeting analyst with expertise in creating comprehensive, strategic meeting reports for C-level executives. You excel at synthesizing complex discussions into actionable insights and clear recommendations. Always respond with valid JSON."
                },
                {"role": "user", "content": prompt}
            ],
            temperature=0.2,  # Lower temperature for more consistent final output
        )
        
        raw_response = completion.choices[0].message.content.strip()
        
        # Parse JSON response
        try:
            if raw_response.startswith("```json"):
                raw_response = raw_response[7:]
            if raw_response.endswith("```"):
                raw_response = raw_response[:-3]
            raw_response = raw_response.strip()
            
            analysis = json.loads(raw_response)
            
            print(f"  ✅ Final meeting report generated successfully")
            return analysis
            
        except json.JSONDecodeError as e:
            print(f"  ⚠️ Failed to parse JSON response for final report: {e}")
            print(f"    Raw response: {raw_response[:200]}...")
            return None
            
    except Exception as e:
        print(f"  ❌ Error generating final report: {e}")
        return None

def save_phase2_results(phase2_results: List[Dict[str, Any]], output_path: str) -> None:
    """Save Phase 2 results to JSON."""
    results = {
        'phase': 2,
        'processing_timestamp': datetime.now().isoformat(),
        'total_meta_chunks': len(phase2_results),
        'meta_chunks': phase2_results,
        'statistics': {
            'total_action_items': sum(len(chunk['phase2_analysis']['consolidated_action_items']) for chunk in phase2_results),
            'total_decisions': sum(len(chunk['phase2_analysis']['important_decisions']) for chunk in phase2_results),
            'total_insights': sum(len(chunk['phase2_analysis']['strategic_insights']) for chunk in phase2_results),
            'total_issues': sum(len(chunk['phase2_analysis']['outstanding_issues']) for chunk in phase2_results)
        }
    }
    
    try:
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(results, f, indent=2, ensure_ascii=False)
        print(f"✅ Phase 2 results saved to: {output_path}")
    except Exception as e:
        print(f"❌ Error saving Phase 2 results: {e}")

def generate_final_docx_report(phase3_results: Dict[str, Any], phase2_results: List[Dict[str, Any]], 
                              original_stats: Dict[str, Any], output_path: str) -> None:
    """Generate comprehensive final DOCX report."""
    try:
        doc = Document()
        
        # Title
        title = doc.add_heading('COMPREHENSIVE MEETING REPORT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle with date
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f"Generated on {datetime.now().strftime('%B %d, %Y')}")
        run.italic = True
        
        # Executive Summary
        doc.add_heading('EXECUTIVE SUMMARY', level=1)
        doc.add_paragraph(phase3_results['executive_summary'])
        
        # Meeting Overview
        doc.add_heading('MEETING OVERVIEW', level=1)
        overview_para = doc.add_paragraph()
        overview_para.add_run('Purpose: ').bold = True
        overview_para.add_run(f"{phase3_results.get('meeting_purpose', 'Strategy and collaboration discussion')}\n")
        overview_para.add_run('Participants: ').bold = True
        overview_para.add_run(f"{', '.join(phase3_results.get('key_participants', ['Multiple speakers']))}\n")
        overview_para.add_run('Total Segments Analyzed: ').bold = True
        overview_para.add_run(f"{original_stats.get('total_segments_processed', 'N/A')}\n")
        overview_para.add_run('Analysis Phases: ').bold = True
        overview_para.add_run("3 (Initial chunking, intermediate synthesis, final consolidation)")
        
        # Strategic Outcomes
        if phase3_results.get('strategic_outcomes'):
            doc.add_heading('STRATEGIC OUTCOMES', level=1)
            for i, outcome in enumerate(phase3_results['strategic_outcomes'], 1):
                outcome_para = doc.add_paragraph()
                outcome_para.add_run(f"{i}. ").bold = True
                outcome_para.add_run(outcome)
        
        # Action Items
        doc.add_heading('ACTION ITEMS', level=1)
        if phase3_results.get('final_action_items'):
            for item in phase3_results['final_action_items']:
                item_para = doc.add_paragraph()
                item_para.add_run(f"#{item.get('id', 'N/A')} - ").bold = True
                item_para.add_run(f"{item['description']}\n")
                item_para.add_run('   Assignee: ').italic = True
                item_para.add_run(f"{item['assignee']}\n")
                item_para.add_run('   Deadline: ').italic = True
                item_para.add_run(f"{item['deadline']}\n")
                item_para.add_run('   Priority: ').italic = True
                item_para.add_run(f"{item.get('priority', 'Not specified')}\n")
                if item.get('dependencies'):
                    item_para.add_run('   Dependencies: ').italic = True
                    item_para.add_run(f"{item['dependencies']}\n")
                if item.get('success_criteria'):
                    item_para.add_run('   Success Criteria: ').italic = True
                    item_para.add_run(f"{item['success_criteria']}")
        else:
            doc.add_paragraph("No specific action items were identified during this meeting.")
        
        # Decisions and Resolutions
        doc.add_heading('DECISIONS AND RESOLUTIONS', level=1)
        if phase3_results.get('decisions_and_resolutions'):
            for i, decision in enumerate(phase3_results['decisions_and_resolutions'], 1):
                decision_para = doc.add_paragraph()
                decision_para.add_run(f"{i}. ").bold = True
                decision_para.add_run(f"{decision['decision']}\n")
                decision_para.add_run('   Rationale: ').italic = True
                decision_para.add_run(f"{decision.get('rationale', 'Not specified')}\n")
                decision_para.add_run('   Impact: ').italic = True
                decision_para.add_run(f"{decision.get('impact', 'To be determined')}")
        else:
            doc.add_paragraph("No major decisions were documented during this meeting.")
        
        # Main Topics Discussed
        doc.add_heading('MAIN TOPICS DISCUSSED', level=1)
        if phase3_results.get('main_topics_discussed'):
            for topic in phase3_results['main_topics_discussed']:
                topic_para = doc.add_paragraph()
                topic_para.add_run(f"{topic['topic']}").bold = True
                topic_para.add_run(f"\n{topic['description']}")
                if topic.get('outcome'):
                    topic_para.add_run(f"\nOutcome: {topic['outcome']}")
        
        # Risks and Concerns
        if phase3_results.get('risks_and_concerns'):
            doc.add_heading('RISKS AND CONCERNS', level=1)
            for i, risk in enumerate(phase3_results['risks_and_concerns'], 1):
                risk_para = doc.add_paragraph()
                risk_para.add_run(f"{i}. ").bold = True
                risk_para.add_run(risk)
        
        # Recommendations
        if phase3_results.get('recommendations'):
            doc.add_heading('RECOMMENDATIONS', level=1)
            for i, recommendation in enumerate(phase3_results['recommendations'], 1):
                rec_para = doc.add_paragraph()
                rec_para.add_run(f"{i}. ").bold = True
                rec_para.add_run(recommendation)
        
        # Follow-up Required
        if phase3_results.get('follow_up_required'):
            doc.add_heading('FOLLOW-UP REQUIRED', level=1)
            for i, followup in enumerate(phase3_results['follow_up_required'], 1):
                followup_para = doc.add_paragraph()
                followup_para.add_run(f"{i}. ").bold = True
                followup_para.add_run(followup)
        
        # Meeting Effectiveness Assessment
        if phase3_results.get('meeting_effectiveness_assessment'):
            doc.add_heading('MEETING EFFECTIVENESS ASSESSMENT', level=1)
            doc.add_paragraph(phase3_results['meeting_effectiveness_assessment'])
        
        # Appendix - Phase 2 Intermediate Analysis
        doc.add_heading('APPENDIX: INTERMEDIATE ANALYSIS DETAILS', level=1)
        for meta_chunk in phase2_results:
            analysis = meta_chunk['phase2_analysis']
            doc.add_heading(f"Section {meta_chunk['meta_chunk_id']}", level=2)
            
            section_para = doc.add_paragraph()
            section_para.add_run('Time Range: ').bold = True
            section_para.add_run(f"{meta_chunk['metadata']['start_time']} - {meta_chunk['metadata']['end_time']}\n")
            section_para.add_run('Summary: ').bold = True
            section_para.add_run(f"{analysis['intermediate_summary']}\n")
            
            if analysis.get('key_themes'):
                section_para.add_run('Key Themes: ').bold = True
                section_para.add_run(f"{', '.join(analysis['key_themes'])}")
        
        # Save document
        doc.save(output_path)
        print(f"✅ Final DOCX report saved to: {output_path}")
        
    except Exception as e:
        print(f"❌ Error generating final DOCX report: {e}")

def main():
    if len(sys.argv) < 2:
        print("Usage: python generate_report_phase2_3.py <phase1_json_file> [output_base_name]")
        print("Example: python generate_report_phase2_3.py summary.json final_report")
        sys.exit(1)
    
    phase1_json_path = sys.argv[1]
    
    if not os.path.exists(phase1_json_path):
        print(f"❌ Phase 1 JSON file not found: {phase1_json_path}")
        sys.exit(1)
    
    # Determine output base name
    if len(sys.argv) > 2:
        output_base = sys.argv[2]
    else:
        base, _ = os.path.splitext(phase1_json_path)
        output_base = f"{base}_final"
    
    phase2_json_path = f"{output_base}_phase2.json"
    phase3_json_path = f"{output_base}_phase3.json"
    final_docx_path = f"{output_base}_meeting_report.docx"
    
    print("🚀 Meeting Report Generator - Phase 2 & 3")
    print(f"   Input Phase 1 JSON: {phase1_json_path}")
    print(f"   Output Phase 2 JSON: {phase2_json_path}")
    print(f"   Output Phase 3 JSON: {phase3_json_path}")
    print(f"   Output Final DOCX: {final_docx_path}")
    print(f"   LLM Endpoint: {LLM_API_BASE_URL}")
    print(f"   LLM Model: {LLM_MODEL_NAME}")
    print(f"   Meta-chunk size: {META_CHUNK_SIZE} Phase 1 chunks")
    
    # Initialize LLM client
    try:
        client = OpenAI(base_url=LLM_API_BASE_URL, api_key=LLM_API_KEY)
        print("✅ LLM client initialized")
    except Exception as e:
        print(f"❌ Failed to initialize LLM client: {e}")
        sys.exit(1)
    
    # Load Phase 1 results
    phase1_data = load_phase1_results(phase1_json_path)
    
    # Phase 2: Create and process meta-chunks
    print(f"\n📦 Creating meta-chunks from {len(phase1_data['chunks'])} Phase 1 chunks...")
    meta_chunks = create_meta_chunks(phase1_data['chunks'], META_CHUNK_SIZE)
    print(f"✅ Created {len(meta_chunks)} meta-chunks")
    
    # Process Phase 2
    phase2_results = process_phase2(meta_chunks, client, LLM_MODEL_NAME)
    
    # Save Phase 2 results
    save_phase2_results(phase2_results, phase2_json_path)
    
    # Phase 3: Generate final report
    print(f"\n🎯 Phase 3: Generating comprehensive final meeting report...")
    phase3_prompt = build_phase3_prompt(phase2_results, phase1_data['statistics'])
    phase3_analysis = query_llm_phase3(client, phase3_prompt, LLM_MODEL_NAME)
    
    if phase3_analysis:
        # Save Phase 3 results
        phase3_results = {
            'phase': 3,
            'processing_timestamp': datetime.now().isoformat(),
            'final_analysis': phase3_analysis,
            'source_meta_chunks': len(phase2_results),
            'original_segments': phase1_data['statistics']['total_segments_processed']
        }
        
        try:
            with open(phase3_json_path, 'w', encoding='utf-8') as f:
                json.dump(phase3_results, f, indent=2, ensure_ascii=False)
            print(f"✅ Phase 3 results saved to: {phase3_json_path}")
        except Exception as e:
            print(f"❌ Error saving Phase 3 results: {e}")
        
        # Generate final DOCX report
        generate_final_docx_report(phase3_analysis, phase2_results, phase1_data['statistics'], final_docx_path)
        
        # Print summary statistics
        print(f"\n🏁 Hierarchical Map-Reduce Analysis Complete!")
        print(f"   📊 Original segments: {phase1_data['statistics']['total_segments_processed']}")
        print(f"   📦 Phase 1 chunks: {len(phase1_data['chunks'])}")
        print(f"   🔄 Phase 2 meta-chunks: {len(phase2_results)}")
        print(f"   🎯 Final action items: {len(phase3_analysis.get('final_action_items', []))}")
        print(f"   ⚖️ Major decisions: {len(phase3_analysis.get('decisions_and_resolutions', []))}")
        print(f"   📄 Final report: {final_docx_path}")
        
    else:
        print("❌ Phase 3 analysis failed. Check LLM connection and response.")
        sys.exit(1)

if __name__ == "__main__":
    main() 